import gspread
from google.oauth2.service_account import Credentials
from docx import Document
import comtypes.client
import os
import time
import re

# Function to fetch data from Google Sheets
def fetch_google_sheet_data(sheet_id, sheet_name, credentials_file):
    try:
        # Define the scope for the Google Sheets API
        scope = [
            'https://www.googleapis.com/auth/spreadsheets',
            'https://www.googleapis.com/auth/drive'
        ]
        
        # Authenticate using the service account credentials
        creds = Credentials.from_service_account_file(credentials_file, scopes=scope)
        client = gspread.authorize(creds)
        
        # Open the Google Sheet and the specific sheet/tab
        sheet = client.open_by_key(sheet_id).worksheet(sheet_name)
        
        # Fetch all data
        data = sheet.get_all_values()
        return data
    
    except Exception as e:
        print(f"An error occurred: {e}")
        return None

# Function to convert DOCX to PDF
def convert_docx_to_pdf(input_path, output_path):
    print(f"Converting file: {input_path} to PDF at {output_path}")
    if not os.path.exists(input_path):
        print(f"File does not exist: {input_path}")
        return
    
    try:
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(input_path)
        doc.SaveAs(output_path, FileFormat=17)  # 17 stands for PDF
        doc.Close()
        word.Quit()
        print(f"PDF successfully created at {output_path}")
    except Exception as e:
        print(f"Error converting to PDF for {input_path}: {e}")

# Function to generate salary slips and convert to PDF
def generate_salary_slip(template_path, output_dir, employee_data, headers):
    # Load the Word template
    template = Document(template_path)
    
    # Map employee data to headers (placeholders)
    placeholder_data = dict(zip(headers, employee_data))
    
    # Replace placeholders with actual data
    for paragraph in template.paragraphs:
        for run in paragraph.runs:
            for placeholder, value in placeholder_data.items():
                if f"{{{placeholder}}}" in run.text:
                    run.text = run.text.replace(f"{{{placeholder}}}", value)
    
    # Replace placeholders in tables
    for table in template.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for placeholder, value in placeholder_data.items():
                            if f"{{{placeholder}}}" in run.text:
                                run.text = run.text.replace(f"{{{placeholder}}}", value)
    
    # Sanitize the employee name for file naming
    employee_name = re.sub(r'[^\w\s]', '', placeholder_data.get("Name", "Employee")).replace(" ", "_")
    
    # Save the updated document locally
    output_docx_path = os.path.join(output_dir, f"{employee_name}_Salary_Slip.docx")
    template.save(output_docx_path)
    
    # Verify the file exists
    if not os.path.exists(output_docx_path):
        print(f"Error: .docx file not found after saving at {output_docx_path}")
        return
    print(f"Salary slip saved for {placeholder_data.get('Name')} at {output_docx_path}")
    
    # Pause briefly to ensure the file is fully saved
    time.sleep(1)
    
    # Convert the saved .docx to .pdf
    output_pdf_path = os.path.join(output_dir, f"{employee_name}_Salary_Slip.pdf")
    convert_docx_to_pdf(output_docx_path, output_pdf_path)

# Main execution
sheet_id = "place your sheet id here"
sheet_name = "Sheet1"  # Replace with the name of your sheet tab
credentials_file = r"credentials.json"
template_path = r"ssformat.docx"  # Path to your Word template
output_dir = r"C:\Users\Kaustubh\Desktop\project\Salary_Slips"  # Directory to save salary slips

# Ensure the output directory exists
os.makedirs(output_dir, exist_ok=True)

# Fetch data from Google Sheet
data = fetch_google_sheet_data(sheet_id, sheet_name, credentials_file)

if data:
    headers = data[0]  # First row contains headers/placeholders
    employees = data[1:]  # Remaining rows contain employee data

    for employee_data in employees:
        generate_salary_slip(template_path, output_dir, employee_data, headers)
else:
    print("No data fetched from the Google Sheet.")
